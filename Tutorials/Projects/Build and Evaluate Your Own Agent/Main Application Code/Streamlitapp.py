import streamlit as st
from streamlit_chat import message
from langchain.tools import tool
from langchain_openai import ChatOpenAI
from langchain.agents import AgentType, initialize_agent
from langchain.prompts.chat import ChatPromptTemplate, MessagesPlaceholder
from langchain.memory import ConversationBufferMemory
import requests
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# Initialize session state with memory
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
    st.session_state.memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
if 'last_output' not in st.session_state:
    st.session_state.last_output = ""

# Email configuration
EMAIL_CONFIG = {
    "smtp_server": "smtp.gmail.com",
    "smtp_port": 587,
    "sender_email": "",
    "sender_password": ""
}

# --- Tool Definitions ---
@tool
def brainstorm_startup_ideas(
    interests: str,
    industry: str = "",
    skills: str = "",
    pain_points: str = ""
) -> dict:
    """Generate creative startup ideas based on interests, industry, skills, and pain points."""
    url = ""
    my_token = ""

    user_prompt = f"""Generate creative and actionable startup or project ideas based on the following:
- Interests: {interests}
- Industry: {industry}
- Skills: {skills}
- Pain Points: {pain_points}
Format the output clearly with bullet points or numbers."""

    data = {
        "input_text": "",
        "system_prompt": "You are a helpful assistant who generates startup ideas tailored to user input.",
        "user_prompt": user_prompt,
        "temperature": 0.7,
        "monitor_model": True,
        "knowledge_base_ids": [],
        "session_id": "",
        "images_urls": []
    }

    try:
        response = requests.post(url + my_token, json=data)
        if response.status_code == 200:
            res = response.json()
            return {"ideas": res.get("output")}
        else:
            return {"error": f"{response.status_code} - {response.text}"}
    except Exception as e:
        return {"error": str(e)}


@tool
def validate_idea(
    idea_description: str,
    target_audience: str = "",
    value_proposition: str = "",
    assumptions: str = ""
) -> dict:
    """Validate a startup idea using lean startup framework and provide validation steps."""
    url = ""
    my_token = ""

    user_prompt = f"""Validate the following idea using a lean startup framework. Provide a checklist or set of validation questions that help assess its feasibility.

Idea: {idea_description}
Target Audience: {target_audience}
Value Proposition: {value_proposition}
Key Assumptions: {assumptions}

Give practical and direct validation steps or criteria.
"""

    data = {
        "input_text": "",
        "system_prompt": "You are a lean startup expert helping users validate ideas realistically.",
        "user_prompt": user_prompt,
        "temperature": 0.7,
        "monitor_model": True,
        "knowledge_base_ids": [],
        "session_id": "",
        "images_urls": []
    }

    try:
        response = requests.post(url + my_token, json=data)
        if response.status_code == 200:
            res = response.json()
            return {"validation": res.get("output")}
        else:
            return {"error": f"{response.status_code} - {response.text}"}
    except Exception as e:
        return {"error": str(e)}


@tool
def generate_lean_canvas(
    idea: str,
    target_market: str = "",
    value_proposition: str = ""
) -> dict:
    """Generate a Lean Canvas for a startup idea with detailed sections."""
    url = ""
    my_token = ""

    user_prompt = f"""Create a Lean Canvas for the following startup idea. Provide clear sections such as Problem, Solution, Key Metrics, Unique Value Proposition, Channels, Customer Segments, Cost Structure, Revenue Streams, and Unfair Advantage.

Idea: {idea}
Target Market: {target_market}
Value Proposition: {value_proposition}

Format the canvas clearly with section headers.
"""

    data = {
        "input_text": "",
        "system_prompt": "You are a business strategist generating detailed Lean Canvas models for startups.",
        "user_prompt": user_prompt,
        "temperature": 0.7,
        "monitor_model": True,
        "knowledge_base_ids": [],
        "session_id": "",
        "images_urls": []
    }

    try:
        response = requests.post(url + my_token, json=data)
        if response.status_code == 200:
            res = response.json()
            return {"lean_canvas": res.get("output")}
        else:
            return {"error": f"{response.status_code} - {response.text}"}
    except Exception as e:
        return {"error": str(e)}


@tool
def save_to_word_doc(content: str, filename: str = "business_canvas.docx") -> dict:
    """Save text content to a Word document."""
    try:
        doc = Document()
        doc.add_heading("Business Canvas Document", 0)
        
        # Use the last output if no specific content provided
        if not content and st.session_state.last_output:
            content = st.session_state.last_output
            
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(filename)
        return {"filename": filename, "message": f"Document saved as {filename}"}
    except Exception as e:
        return {"error": str(e)}


@tool
def send_email(
    recipient_email: str,
    subject: str,
    body: str,
    sender_email: str = None,
    sender_password: str = None
) -> dict:
    """Send an email using SMTP."""
    try:
        from_email = sender_email if sender_email else EMAIL_CONFIG["sender_email"]
        password = sender_password if sender_password else EMAIL_CONFIG["sender_password"]

        if not from_email or not password:
            return {"error": "Email credentials not configured"}

        msg = MIMEMultipart()
        msg['From'] = from_email
        msg['To'] = recipient_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain'))

        server = smtplib.SMTP(EMAIL_CONFIG["smtp_server"], EMAIL_CONFIG["smtp_port"])
        server.starttls()
        server.login(from_email, password)
        server.sendmail(from_email, recipient_email, msg.as_string())
        server.quit()

        return {"status": "success", "message": f"Email sent to {recipient_email}"}
    except Exception as e:
        return {"error": str(e)}

# Initialize the agent with memory
def initialize_agent_executor():
    tools = [
        brainstorm_startup_ideas,
        validate_idea,
        generate_lean_canvas,
        save_to_word_doc,
        send_email
    ]

    llm = ChatOpenAI(model="gpt-4", openai_api_key="")

    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", """You are a helpful startup strategist assistant. Remember:
            1. Maintain context throughout conversations
            2. Reference previous outputs when appropriate
            3. For save operations, use the most recent generated content unless specified otherwise
            4. Be specific when asking for clarifications"""),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}"),
            MessagesPlaceholder(variable_name="agent_scratchpad"),
        ]
    )

    return initialize_agent(
        tools,
        llm,
        agent=AgentType.OPENAI_FUNCTIONS,
        verbose=True,
        memory=st.session_state.memory,
        prompt=prompt
    )

# --- Streamlit UI ---
st.set_page_config(page_title="🚀 Startup Strategist", page_icon="🚀")

# Sidebar configuration
with st.sidebar:
    st.title("Configuration")
    st.subheader("Email Settings")
    
    EMAIL_CONFIG["sender_email"] = st.text_input(
        "Sender Email", 
        value=EMAIL_CONFIG["sender_email"],
        help="Your email address for sending emails"
    )
    
    EMAIL_CONFIG["sender_password"] = st.text_input(
        "Sender Password", 
        value=EMAIL_CONFIG["sender_password"],
        type="password",
        help="Your email password or app password"
    )
    
    st.markdown("---")
    st.write("**Available Functions:**")
    st.write("• Brainstorm startup ideas")
    st.write("• Validate business concepts")
    st.write("• Generate lean canvases")
    st.write("• Save documents to Word")
    st.write("• Send emails")

# Main chat interface
st.title("🚀 Startup Strategist Chatbot")
st.caption("Your AI assistant for startup ideation and validation")

# Display chat history
for i, chat in enumerate(st.session_state.chat_history):
    if chat['role'] == 'user':
        message(chat['content'], is_user=True, key=f"user_{i}")
    else:
        message(chat['content'], key=f"bot_{i}")

# User input
user_input = st.chat_input("Type your message here...")

if user_input:
    # Initialize agent if not already done
    if 'agent' not in st.session_state:
        st.session_state.agent = initialize_agent_executor()
    
    # Add user message to chat history
    st.session_state.chat_history.append({"role": "user", "content": user_input})
    
    # Get agent response
    with st.spinner("Thinking..."):
        try:
            response = st.session_state.agent({"input": user_input})
            output = response["output"]
            
            # Store the last output for reference
            st.session_state.last_output = output
            st.session_state.chat_history.append({"role": "assistant", "content": output})
            
        except Exception as e:
            error_msg = f"Sorry, I encountered an error: {str(e)}"
            st.session_state.chat_history.append({"role": "assistant", "content": error_msg})
    
    # Rerun to update the display
    st.rerun()
